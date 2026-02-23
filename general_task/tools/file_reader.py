"""File reader tool for DOCX, JSON, JSONLD, XML, TXT formats."""

from __future__ import annotations

import json
import os
from pathlib import Path

from google.adk.tools import FunctionTool


def file_reader(file_path: str, **kwargs: object) -> str:
    """Read various file formats.

    - .docx: Use python-docx, iterate paragraphs AND tables
    - .json/.jsonld/.xml/.txt: Read as text

    Args:
        file_path: Path to the file
        **kwargs: Extra arguments (e.g., 'page') are ignored
    """
    if not os.path.exists(file_path):
        return f"[ERROR] File not found: {file_path}"

    ext = Path(file_path).suffix.lower()

    try:
        # DOCX files
        if ext == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "[ERROR] python-docx not installed. Run: pip install python-docx"

            doc = Document(file_path)
            output = [f"Document: {os.path.basename(file_path)}\n"]

            # Extract paragraphs
            output.append("=== Paragraphs ===")
            for para in doc.paragraphs:
                if para.text.strip():
                    output.append(f"{para.text}")

            # Extract tables
            if doc.tables:
                output.append(f"\n=== Tables ({len(doc.tables)}) ===")
                for t_idx, table in enumerate(doc.tables, 1):
                    output.append(f"\n--- Table {t_idx} ---")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        output.append(row_text)

            return "\n".join(output)

        # JSON/JSONLD files
        elif ext in [".json", ".jsonld"]:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            try:
                data = json.loads(content)
                return (
                    f"JSON content ({os.path.basename(file_path)}):\n"
                    f"{json.dumps(data, indent=2, ensure_ascii=False)}"
                )
            except json.JSONDecodeError:
                return f"Raw content ({os.path.basename(file_path)}):\n{content}"

        # XML files
        elif ext == ".xml":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return f"XML content ({os.path.basename(file_path)}):\n{content}"

        # TXT files
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return f"Text content ({os.path.basename(file_path)}):\n{content}"

        else:
            return (
                f"[ERROR] Unsupported file format: {ext}."
                " Supported: .docx, .json, .jsonld, .xml, .txt"
            )

    except Exception as e:
        return f"[ERROR] Failed to read file: {str(e)}"


file_reader_tool = FunctionTool(file_reader)
