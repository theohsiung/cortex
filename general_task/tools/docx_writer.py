"""DOCX writer tool for generating Word documents."""

from __future__ import annotations

import os
from typing import Optional

from google.adk.tools import FunctionTool

from .._config import agent_config


def docx_writer(
    file_name: str,
    content: str,
    title: Optional[str] = None,
) -> str:
    """Create a Word (.docx) document from structured content.

    The content uses a simple markup format:
    - Lines starting with '# '   → Heading level 1
    - Lines starting with '## '  → Heading level 2
    - Lines starting with '### ' → Heading level 3
    - Lines starting with '- '   → Bullet list item
    - Lines starting with '1. '  → Numbered list item (any digit)
    - Lines starting with '| '   → Table row (cells separated by ' | ')
    - Lines starting with '---'  → Page break
    - All other lines            → Normal paragraph
    - Empty lines are preserved as spacing

    Args:
        file_name: Output file name (e.g., 'report.docx'). Saved to output directory.
        content: Document content in the markup format described above.
        title: Optional document title added as a Title style heading at the top.
    """
    try:
        from docx import Document
    except ImportError:
        return "[ERROR] python-docx not installed. Run: pip install python-docx"

    try:
        doc = Document()

        if title:
            doc.add_heading(title, level=0)

        # Track table state
        table_rows: list[list[str]] = []

        def _flush_table() -> None:
            """Write accumulated table rows into the document."""
            if not table_rows:
                return
            num_cols = max(len(row) for row in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = "Table Grid"
            for r_idx, row_data in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        table.rows[r_idx].cells[c_idx].text = cell_text
            table_rows.clear()

        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()

            # Table row detection
            if stripped.startswith("| ") and stripped.endswith(" |"):
                inner = stripped[1:-1]
                # Skip separator rows like |---|---|
                if all(c in "-| " for c in inner):
                    continue
                cells = [c.strip() for c in inner.split("|")]
                table_rows.append(cells)
                continue

            # Flush any pending table before non-table content
            _flush_table()

            # Page break
            if stripped.startswith("---"):
                doc.add_page_break()
                continue

            # Empty line
            if not stripped:
                doc.add_paragraph("")
                continue

            # Headings
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            # Bullet list
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            # Numbered list
            elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:5]:
                text = stripped[stripped.index(". ") + 2 :]
                doc.add_paragraph(text, style="List Number")
            # Normal paragraph
            else:
                doc.add_paragraph(stripped)

        # Flush remaining table
        _flush_table()

        # Save
        output_dir = agent_config.output_dir
        os.makedirs(output_dir, exist_ok=True)
        if not file_name.endswith(".docx"):
            file_name += ".docx"
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)

        return f"[SUCCESS] Document saved to: {output_path}"

    except Exception as e:
        return f"[ERROR] Failed to create document: {str(e)}"


docx_writer_tool = FunctionTool(docx_writer)
